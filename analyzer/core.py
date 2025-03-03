# analyzer/core.py
import re
import spacy
import pandas as pd
import os
from typing import List, Dict, Any, Tuple
import docx
from PyPDF2 import PdfReader
import logging
from pathlib import Path
from .validators import (
    validate_email, validate_phone, validate_date, 
    validate_secu, validate_siret, validate_person_name,
    validate_postal_address, validate_ip_address
)

# Chargement du modèle spaCy (sera initialisé au premier appel)
nlp = None

# Listes d'exclusion importées du module de configuration
from config.exclusion_lists import (
    EXCLUDED_PERSONS, PROFESSIONAL_CONTEXT, 
    TEMPLATE_INDICATORS, ORGANIZATION_UNITS
)

# Expressions régulières améliorées
EMAIL_REGEX = re.compile(
    r'\b[a-zA-Z0-9][a-zA-Z0-9._%+-]{0,63}@(?:[a-zA-Z0-9-]{1,63}\.){1,8}[a-zA-Z]{2,63}\b'
)
PHONE_REGEX = re.compile(
    r'\b(?:(?:(?:\+|00)33[ .-]?(?:\(0\)[ .-]?)?)|0)[1-9](?:[ .-]?\d{2}){4}\b'
)
DATE_REGEX = re.compile(
    r'\b(?:0[1-9]|[12]\d|3[01])[-/.](0[1-9]|1[012])[-/.](19|20)\d{2}\b'
)
SECU_REGEX = re.compile(
    r'\b[123]\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01]|[2468][02468]|[13579][13579])\d{6}(?:\d{2})?\b'
)
SIRET_REGEX = re.compile(r'\b\d{14}\b')

# ===========================
# Expressions régulières améliorées pour adresses postales et IP
# ===========================
POSTAL_ADDRESS_REGEX = re.compile(
    r'\b\d{1,4}[,\s]+(?:[a-zA-ZÀ-ÿ\'\-\.\s]+)[,\s]+\d{5}(?:\s+[a-zA-ZÀ-ÿ\'\-\.\s]+)?\b'
)
IP_ADDRESS_REGEX = re.compile(
    r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b'
)

def initialize_nlp():
    """Initialise le modèle NLP et les patterns personnalisés."""
    global nlp
    if nlp is None:
        try:
            SPACY_MODEL = "fr_core_news_md"
            nlp = spacy.load(SPACY_MODEL)
            add_custom_patterns(nlp)
        except OSError:
            logging.error(f"Modèle spaCy {SPACY_MODEL} non trouvé. Installation nécessaire.")
            raise Exception(f"Modèle spaCy non trouvé. Exécutez : python -m spacy download {SPACY_MODEL}")
    return nlp

def add_custom_patterns(nlp):
    """
    Ajoute des patterns personnalisés au modèle NER pour exclure les entités spécifiques.
    Améliore la détection en distinguant personnes vs organisations.
    """
    try:
        # Si le ruler existe déjà, on le récupère, sinon on le crée
        if "entity_ruler" in nlp.pipe_names:
            ruler = nlp.get_pipe("entity_ruler")
        else:
            ruler = nlp.add_pipe("entity_ruler", before="ner")
        
        # Patterns pour ignorer les noms spécifiques
        patterns = []
        for person in EXCLUDED_PERSONS:
            patterns.append({"label": "IGNORED_ENTITY", "pattern": person})
        
        # Patterns pour ignorer les unités organisationnelles
        for org_unit in ORGANIZATION_UNITS:
            patterns.append({"label": "ORG", "pattern": org_unit})
        
        # Patterns pour ignorer les formats standard souvent présents dans les templates
        template_patterns = [
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "nom"}, {"LOWER": "prénom"}]},
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "prénom"}, {"LOWER": "nom"}]},
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "mr"}, {"LOWER": "x"}]},
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "monsieur"}, {"LOWER": "x"}]},
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "madame"}, {"LOWER": "x"}]},
            {"label": "IGNORED_ENTITY", "pattern": [{"LOWER": "mme"}, {"LOWER": "x"}]},
            {"label": "IGNORED_ENTITY", "pattern": "John Doe"},
            {"label": "IGNORED_ENTITY", "pattern": "Jane Doe"},
        ]
        patterns.extend(template_patterns)
        
        ruler.add_patterns(patterns)
        logging.info(f"Ajout de {len(patterns)} patterns pour exclure des entités spécifiques")
        return True
    except Exception as e:
        logging.error(f"Erreur lors de l'ajout de patterns personnalisés: {str(e)}")
        return False

def is_likely_organizational_name(text: str, entity: str) -> bool:
    """
    Détermine si un nom est probablement lié à l'organisation plutôt qu'à une personne externe.
    Analyse détaillée du contexte pour réduire les faux positifs.
    """
    text_lower = text.lower()
    entity_lower = entity.lower()
    
    # Vérifier si l'entité contient une unité organisationnelle connue
    if any(unit.lower() in entity_lower for unit in ORGANIZATION_UNITS):
        return True
    
    # Vérifier si l'entité est complètement en majuscules (acronyme)
    if entity.isupper() and len(entity.split()) <= 2:
        return True
    
    # Vérifier les indicateurs d'organisation dans le nom lui-même
    org_name_indicators = ["service", "département", "direction", "pôle", "équipe", "groupe", "unité"]
    if any(indicator in entity_lower for indicator in org_name_indicators):
        return True
        
    # Indicateurs de contexte professionnel autour de l'entité
    org_indicators = [
        "directeur", "directrice", "responsable", "chef", "technicien", "informatique",
        "référent", "chargé de", "service", "département", "pôle", "l'équipe",
        "signature", "contact", "coordonnées", "adjoint", "administratif",
        "conseiller", "manager", "gestion", "gestionnaire", "assistant"
    ]
    
    # Analyse du contexte proximal
    occurrences = []
    start_idx = 0
    while True:
        idx = text_lower.find(entity_lower, start_idx)
        if idx == -1:
            break
        window_start = max(0, idx - 50)  # Fenêtre plus large
        window_end = min(len(text), idx + len(entity) + 50)
        context = text[window_start:window_end].lower()
        occurrences.append(context)
        start_idx = idx + 1
    
    # Si aucune occurrence, retourner False
    if not occurrences:
        return False
    
    # Compter les contextes organisationnels
    org_contexts = sum(1 for context in occurrences if any(indicator in context for indicator in org_indicators))
    
    # Si plus de la moitié des occurrences sont dans un contexte organisationnel
    if org_contexts >= len(occurrences) / 2:
        return True
    
    # Nouveaux patterns spécifiques aux formules officielles
    official_patterns = [
        f"m. {entity_lower}", f"mr {entity_lower}", f"mme {entity_lower}",
        f"monsieur {entity_lower}", f"madame {entity_lower}",
        f"{entity_lower}, directeur", f"{entity_lower}, responsable",
        f"{entity_lower} (directeur", f"{entity_lower} (responsable",
        f"{entity_lower} - directeur", f"{entity_lower} - responsable"
    ]
    
    if any(pattern in text_lower for pattern in official_patterns):
        return True
    
    # Vérifier les mots après l'entité qui indiquent un rôle organisationnel
    for occurrence in occurrences:
        entity_pos = occurrence.find(entity_lower)
        if entity_pos != -1:
            after_entity = occurrence[entity_pos + len(entity_lower):].strip()
            if after_entity.startswith(("est", "a été nommé", "occupe", "en charge", ":")):
                return True
    
    return False

def analyze_name_context(name: str, text: str) -> float:
    """
    Analyse le contexte autour d'un nom pour déterminer s'il s'agit d'un contexte professionnel.
    
    Returns:
        float: Score de contexte professionnel (plus élevé = plus professionnel)
    """
    context_score = 0.0
    
    try:
        name_pos = text.lower().find(name.lower())
        if name_pos == -1:
            return 0.0
            
        start = max(0, name_pos - 100)
        end = min(len(text), name_pos + len(name) + 100)
        context = text[start:end].lower()
        
        for term in PROFESSIONAL_CONTEXT:
            if term in context:
                context_score += 0.15
        
        titles = ["m.", "mme.", "mr.", "dr.", "monsieur", "madame", "docteur", "prof.", "professeur"]
        for title in titles:
            if title in context:
                context_score += 0.1
        
        for indicator in TEMPLATE_INDICATORS:
            if indicator in context:
                context_score += 0.2
    except:
        pass
    
    return min(1.0, context_score)

def is_supported_file(file_path: str) -> bool:
    """Vérifie si le fichier est d'un type supporté pour l'analyse."""
    from .file_utils import is_temp_file
    
    # Ignorer les fichiers temporaires
    if is_temp_file(file_path):
        return False
    
    # Vérifier l'extension du fichier
    extension = Path(file_path).suffix.lower()
    return extension in ['.txt', '.log', '.csv', '.docx', '.doc', '.xlsx', '.xls', '.pdf', '.rtf', '.odt', '.ods']

def get_file_type(file_path: str) -> str:
    """Détermine le type de fichier."""
    extension = Path(file_path).suffix.lower()
    if extension in ['.txt', '.log', '.csv']:
        return 'text'
    elif extension in ['.docx', '.doc']:
        return 'word'
    elif extension in ['.xlsx', '.xls']:
        return 'excel'
    elif extension == '.pdf':
        return 'pdf'
    else:
        return 'unknown'

def read_txt_file(file_path: str) -> str:
    """Lit un fichier texte avec gestion des encodages et des erreurs."""
    from .file_utils import ensure_readable, fix_network_path
    
    # Vérifier que le fichier est accessible et lisible
    if not ensure_readable(file_path):
        return ""
        
    # Corriger les chemins réseau si nécessaire
    file_path = fix_network_path(file_path)
    
    # Essayer différents encodages
    encodings = ['utf-8', 'latin1', 'cp1252', 'utf-16', 'ascii']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
                # Vérifier si le contenu est valide
                if content and not content.isspace():
                    return content
        except UnicodeDecodeError:
            continue
        except PermissionError:
            logging.warning(f"Permission refusée pour {file_path}")
            return ""
        except OSError as e:
            logging.warning(f"Erreur OS lors de la lecture de {file_path}: {str(e)}")
            return ""
    
    logging.warning(f"Impossible de lire {file_path} avec les encodages standards")
    return ""

def read_docx_file(file_path: str) -> str:
    """Lit un fichier DOCX avec gestion des erreurs améliorée."""
    from .file_utils import ensure_readable, is_temp_file, fix_network_path
    
    # Ignorer les fichiers temporaires de Word (commençant par ~$)
    if is_temp_file(file_path):
        logging.info(f"Fichier temporaire Word ignoré: {file_path}")
        return ""
    
    # Vérifier que le fichier est accessible et lisible
    if not ensure_readable(file_path):
        return ""
        
    # Corriger les chemins réseau si nécessaire
    file_path = fix_network_path(file_path)    
    
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        # Si le texte est vide, essayer de vérifier les tableaux
        if not text or text.isspace():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        return text
    except (docx.opc.exceptions.PackageNotFoundError, ValueError):
        logging.info(f"Fichier DOCX inaccessible ou corrompu: {file_path}")
        return ""
    except Exception as e:
        logging.error(f"Erreur lecture DOCX {file_path}: {str(e)}")
        return ""

def read_excel_file(file_path: str) -> str:
    """Lit un fichier Excel avec gestion des erreurs améliorée."""
    from .file_utils import ensure_readable, is_temp_file, fix_network_path
    
    # Ignorer les fichiers temporaires d'Excel
    if is_temp_file(file_path):
        logging.info(f"Fichier temporaire Excel ignoré: {file_path}")
        return ""
    
    # Vérifier que le fichier est accessible et lisible
    if not ensure_readable(file_path):
        return ""
        
    # Corriger les chemins réseau si nécessaire
    file_path = fix_network_path(file_path)
    
    try:
        # Vérifier l'extension pour utiliser le bon moteur
        ext = Path(file_path).suffix.lower()
        if ext == '.xls':
            # Anciens fichiers Excel (.xls)
            try:
                import xlrd
                with xlrd.open_workbook(file_path) as wb:
                    sheet = wb.sheet_by_index(0)
                    text_content = []
                    for i in range(sheet.nrows):
                        row = [str(sheet.cell_value(i, j)) for j in range(sheet.ncols)]
                        text_content.append(' '.join(row))
                    return '\n'.join(text_content)
            except ImportError:
                logging.warning(f"Module xlrd non disponible pour lire {file_path}. Installation: pip install xlrd>=2.0.1")
                return ""
        else:
            # Fichiers Excel plus récents (.xlsx, .xlsm)
            df = pd.read_excel(file_path, engine='openpyxl')
            text_content = []
            for column in df.columns:
                # Convertir la colonne en chaîne avec gestion des erreurs
                try:
                    values = df[column].fillna('').astype(str)
                    text_content.append(f"{column}: {' '.join(values)}")
                except Exception as e:
                    text_content.append(f"{column}: [Erreur de conversion]")
            return "\n".join(text_content)
    except Exception as e:
        logging.error(f"Erreur lecture Excel {file_path}: {str(e)}")
        return ""

def read_pdf_file(file_path: str) -> str:
    """Lit un fichier PDF avec gestion robuste des erreurs."""
    from .file_utils import ensure_readable, fix_network_path
    
    # Vérifier que le fichier est accessible et lisible
    if not ensure_readable(file_path):
        return ""
        
    # Corriger les chemins réseau si nécessaire
    file_path = fix_network_path(file_path)
    
    try:
        # Essayer d'abord avec PyPDF2
        with open(file_path, 'rb') as file:
            try:
                reader = PdfReader(file, strict=False)
                text = []
                for page in reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except Exception as page_e:
                        logging.warning(f"Erreur d'extraction de page dans {file_path}: {str(page_e)}")
                        continue
                return "\n".join(text)
            except Exception as e:
                logging.warning(f"PyPDF2 a échoué pour {file_path}, erreur: {str(e)}")
                
                # Tenter avec pdfplumber comme alternative si PyPDF2 échoue
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text = []
                        for page in pdf.pages:
                            try:
                                page_text = page.extract_text()
                                if page_text:
                                    text.append(page_text)
                            except Exception:
                                continue
                        return "\n".join(text)
                except (ImportError, Exception) as plumb_e:
                    logging.error(f"Toutes les méthodes d'extraction PDF ont échoué pour {file_path}: {str(plumb_e)}")
                    return ""
    except Exception as e:
        logging.error(f"Erreur lecture PDF {file_path}: {str(e)}")
        return ""

def analyze_file(file_path: str) -> Dict[str, Any]:
    """Analyse un fichier et retourne les résultats avec niveaux de confiance."""
    from .file_utils import should_skip_file
    
    # Vérifier si le fichier doit être ignoré
    if should_skip_file(file_path):
        return None
    
    file_type = get_file_type(file_path)
    text_content = ""

    try:
        if file_type == 'text':
            text_content = read_txt_file(file_path)
        elif file_type == 'word':
            text_content = read_docx_file(file_path)
        elif file_type == 'excel':
            text_content = read_excel_file(file_path)
        elif file_type == 'pdf':
            text_content = read_pdf_file(file_path)

        if text_content:
            personal_data = detect_personal_data(text_content, file_path)
            
            result = {
                "file_path": file_path,
                "file_type": file_type,
                "text_snippet": text_content[:100],
            }
            
            # Pour chaque type de données, extraire les valeurs et les niveaux de confiance
            for data_type in personal_data:
                values = [item["value"] for item in personal_data[data_type]]
                confidences = [item["confidence"] for item in personal_data[data_type]]
                result[f"{data_type}_found"] = ", ".join(values)
                result[f"{data_type}_confidence"] = ", ".join([f"{conf:.2f}" for conf in confidences])
                
                # Définir un facteur de risque (plus élevé pour les données très sensibles)
                risk_factor = 3 if data_type in ["secu", "emails", "phones"] else 1
                if data_type in ["postal_addresses", "ip_addresses"]:
                    risk_factor = 2
                risk_scores = [conf * risk_factor for conf in confidences]
                result[f"{data_type}_risk"] = sum(risk_scores) / len(risk_scores) if risk_scores else 0
            
            return result
    except Exception as e:
        logging.error(f"Erreur analyse fichier {file_path}: {str(e)}")
    
    return None

def calculate_risk_scores(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calcule des scores de risque agrégés pour les résultats."""
    risk_analysis = {
        "high_risk_files": [],
        "medium_risk_files": [],
        "low_risk_files": [],
        "risk_by_type": {
            "emails": 0,
            "phones": 0,
            "names": 0,
            "secu": 0,
            "siret": 0,
            "postal_addresses": 0,
            "ip_addresses": 0
        },
        "top_risky_extensions": {},
        "total_analyzed": len(results)
    }
    
    risk_weights = {
        "secu": 10,
        "emails": 5,
        "phones": 5,
        "names": 3,
        "siret": 2,
        "postal_addresses": 2,
        "ip_addresses": 2
    }
    
    for result in results:
        file_risk = 0
        for data_type in ["emails", "phones", "names", "secu", "siret", "postal_addresses", "ip_addresses"]:
            risk_key = f"{data_type}_risk"
            if risk_key in result and result[risk_key]:
                risk_score = float(result[risk_key]) * risk_weights[data_type]
                file_risk += risk_score
                risk_analysis["risk_by_type"][data_type] += risk_score
        
        if file_risk > 20:
            risk_analysis["high_risk_files"].append({
                "path": result["file_path"],
                "score": file_risk,
                "type": result["file_type"]
            })
        elif file_risk > 10:
            risk_analysis["medium_risk_files"].append({
                "path": result["file_path"],
                "score": file_risk,
                "type": result["file_type"]
            })
        elif file_risk > 0:
            risk_analysis["low_risk_files"].append({
                "path": result["file_path"],
                "score": file_risk,
                "type": result["file_type"]
            })
            
        extension = Path(result["file_path"]).suffix.lower()
        if file_risk > 0:
            risk_analysis["top_risky_extensions"][extension] = risk_analysis["top_risky_extensions"].get(extension, 0) + 1
    
    risk_analysis["high_risk_files"] = sorted(risk_analysis["high_risk_files"], key=lambda x: x["score"], reverse=True)
    risk_analysis["medium_risk_files"] = sorted(risk_analysis["medium_risk_files"], key=lambda x: x["score"], reverse=True)
    risk_analysis["total_high_risk"] = len(risk_analysis["high_risk_files"])
    risk_analysis["total_medium_risk"] = len(risk_analysis["medium_risk_files"])
    risk_analysis["total_low_risk"] = len(risk_analysis["low_risk_files"])
    risk_analysis["top_risky_extensions"] = sorted(risk_analysis["top_risky_extensions"].items(), key=lambda x: x[1], reverse=True)
    
    return risk_analysis

def generate_false_positives_report(results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Génère un rapport des potentiels faux positifs."""
    false_positives = []
    
    for result in results:
        file_path = result["file_path"]
        
        if "names_found" in result and result["names_found"]:
            names = result["names_found"].split(", ")
            confidences = [float(conf) for conf in result["names_confidence"].split(", ")] if "names_confidence" in result and result["names_confidence"] else []
            if len(names) != len(confidences):
                confidences = [0.5] * len(names)
            for i, name in enumerate(names):
                confidence = confidences[i] if i < len(confidences) else 0.5
                is_likely_false_positive = False
                if confidence < 0.4 or len(name.split()) < 2:
                    is_likely_false_positive = True
                if any(org.lower() in name.lower() for org in ORGANIZATION_UNITS):
                    is_likely_false_positive = True
                title_indicators = ["service", "département", "direction", "unité", "pôle", "responsable", "chef"]
                if any(indicator.lower() in name.lower() for indicator in title_indicators):
                    is_likely_false_positive = True
                if is_likely_false_positive:
                    false_positives.append({
                        "file_path": file_path,
                        "data_type": "name",
                        "value": name,
                        "confidence": confidence,
                        "reason": "Nom potentiellement non personnel"
                    })
    
    return false_positives

def detect_personal_data(text: str, file_path: str = "") -> Dict[str, List[Any]]:
    """
    Détecte les données personnelles avec validation et scoring.
    Inclut les détections pour adresses postales et adresses IP.
    
    Returns:
        Dict avec "emails", "phones", "dates", "names", "secu", "siret",
        "postal_addresses" et "ip_addresses" et leurs niveaux de confiance.
    """
    initialize_nlp()
    
    results = {
        "emails": [],
        "phones": [],
        "dates": [],
        "names": [],
        "secu": [],
        "siret": [],
        "postal_addresses": [],
        "ip_addresses": []
    }

    if not text or len(text) < 3:
        return results
        
    # Détection si le document est un template/exemple
    is_template = any(indicator in text.lower() for indicator in TEMPLATE_INDICATORS)
    if is_template:
        logging.info(f"Document détecté comme template/exemple: {file_path}")

    try:
        # Emails
        found_emails = EMAIL_REGEX.findall(text)
        for email in found_emails:
            if validate_email(email):
                # Réduire le score de confiance pour les emails dans des templates
                confidence = 0.7 if is_template else 0.9
                if not any(domain.lower() in email.lower() for domain in ["acme.com", "acme.net", "acme.org", "acme.fr"]):
                    results["emails"].append({
                        "value": email,
                        "confidence": confidence
                    })

        # Téléphones - avec gestion renforcée des formats
        found_phones = PHONE_REGEX.findall(text)
        for phone in found_phones:
            if validate_phone(phone):
                # Réduire la confiance si format standard ou dans un template
                if is_template:
                    confidence = 0.7
                elif re.match(r'^0[1-9](?: \d{2}){4}$', phone):  # Format très courant
                    confidence = 0.75
                else:
                    confidence = 0.85
                results["phones"].append({
                    "value": phone,
                    "confidence": confidence
                })

        # Dates
        found_dates = DATE_REGEX.findall(text)
        if isinstance(found_dates, list) and found_dates and isinstance(found_dates[0], tuple):
            dates = ["/".join(d) for d in found_dates]
        else:
            dates = found_dates if isinstance(found_dates, list) else []
        
        for date in dates:
            if validate_date(date):
                results["dates"].append({
                    "value": date,
                    "confidence": 0.5  # Les dates sont moins confidentielles
                })

        # Numéros de sécurité sociale
        found_secu = SECU_REGEX.findall(text)
        for secu in found_secu:
            if validate_secu(secu):
                # Score très élevé pour ce type de données très sensibles
                results["secu"].append({
                    "value": secu,
                    "confidence": 0.98
                })

        # SIRET
        found_siret = SIRET_REGEX.findall(text)
        for siret in found_siret:
            if validate_siret(siret):
                # Score élevé mais un peu moins que sécu
                results["siret"].append({
                    "value": siret,
                    "confidence": 0.92
                })

        # Détection via spaCy pour les noms avec gestion améliorée du contexte
        try:
            # Limiter la taille du texte pour éviter les problèmes de mémoire avec spaCy
            max_length = min(len(text), 100000)
            doc = nlp(text[:max_length])
            
            for ent in doc.ents:
                # Vérifier si c'est une entité de type personne et qu'elle n'est pas déjà ignorée
                if ent.label_ == "PER" and ent.label_ != "IGNORED_ENTITY":
                    name = ent.text.strip()
                    is_valid, confidence = validate_person_name(name, text)
                    if is_valid and not is_likely_organizational_name(text, name):
                        results["names"].append({
                            "value": name,
                            "confidence": confidence
                        })
        except Exception as e:
            logging.error(f"Erreur lors de l'analyse NER: {str(e)}")

        # Détection d'adresses postales
        found_postal_addresses = POSTAL_ADDRESS_REGEX.findall(text)
        for address in found_postal_addresses:
            if validate_postal_address(address):
                # On réduit la confiance si le document est un template
                confidence = 0.65 if is_template else 0.75
                results["postal_addresses"].append({
                    "value": address,
                    "confidence": confidence
                })

        # Détection d'adresses IP
        found_ips = IP_ADDRESS_REGEX.findall(text)
        for ip in found_ips:
            if validate_ip_address(ip):
                # Ajuster la confiance pour les IPs - plus élevée pour IPs privées
                confidence = 0.85
                # Les IPs de test/locales ont une confiance moindre
                if ip.startswith(("127.", "192.168.", "10.", "172.")):
                    confidence = 0.75
                results["ip_addresses"].append({
                    "value": ip,
                    "confidence": confidence
                })
    except Exception as e:
        logging.error(f"Erreur lors de la détection des données personnelles: {str(e)}")

    # Filtrage final selon seuils de confiance
    filtered_results = {key: [] for key in results}
    
    confidence_thresholds = {
        "emails": 0.7,
        "phones": 0.7,
        "dates": 0.5,
        "names": 0.4,
        "secu": 0.8,
        "siret": 0.8,
        "postal_addresses": 0.7,
        "ip_addresses": 0.7
    }
    
    for data_type, items in results.items():
        threshold = confidence_thresholds[data_type]
        for item in items:
            if item["confidence"] >= threshold:
                filtered_results[data_type].append(item)

    return filtered_results